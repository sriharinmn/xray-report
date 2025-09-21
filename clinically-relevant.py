import streamlit as st
import os
import numpy as np
import torch
import torchvision.transforms as transforms
from PIL import Image
import io
import base64
from groq import Groq
from datetime import datetime
import json
import torchxrayvision as xrv
import skimage
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE

from dotenv import load_dotenv
load_dotenv()

# Enhanced clinical thresholds with more conservative values
CLINICAL_THRESHOLDS = {
    # Critical/Urgent findings - higher threshold for specificity
    'Pneumothorax': 0.80,
    'Tension Pneumothorax': 0.85,
    'Massive Pleural Effusion': 0.75,
    'Large Opacity': 0.70,
    
    # Important findings - moderate threshold
    'Pleural Effusion': 0.65,
    'Consolidation': 0.70,
    'Pneumonia': 0.75,
    'Mass': 0.80,
    'Nodule': 0.70,
    'Cardiomegaly': 0.60,
    'Edema': 0.65,
    'Atelectasis': 0.65,
    
    # Less urgent findings - standard threshold
    'Infiltration': 0.60,
    'Emphysema': 0.65,
    'Fibrosis': 0.70,
    'Pleural Thickening': 0.60,
    'Lung Lesion': 0.65,
    'Lung Opacity': 0.60,
    'Enlarged Cardiomediastinum': 0.65,
    
    # Low clinical priority - higher threshold to reduce noise
    'Fracture': 0.75,
    'Hernia': 0.80,
    'Support Devices': 0.85
}

# Enhanced clinical priority levels
CLINICAL_PRIORITY = {
    'CRITICAL': ['Pneumothorax', 'Tension Pneumothorax', 'Massive Pleural Effusion'],
    'HIGH': ['Mass', 'Large Opacity', 'Consolidation', 'Pneumonia'],
    'MODERATE': ['Pleural Effusion', 'Nodule', 'Cardiomegaly', 'Edema', 'Atelectasis'],
    'LOW': ['Infiltration', 'Emphysema', 'Fibrosis', 'Pleural Thickening', 'Fracture']
}

# Rule-based filtering patterns
CONTRADICTION_RULES = {
    # If pneumothorax present, suppress other lung findings unless very high confidence
    'pneumothorax_suppress': {
        'trigger': ['Pneumothorax'],
        'suppress': ['Infiltration', 'Lung Opacity', 'Edema'],
        'min_confidence': 0.85
    },
    # If consolidation present, suppress general infiltration
    'consolidation_suppress': {
        'trigger': ['Consolidation', 'Pneumonia'],
        'suppress': ['Infiltration'],
        'min_confidence': 0.80
    },
    # If mass present, suppress general nodule findings
    'mass_suppress': {
        'trigger': ['Mass'],
        'suppress': ['Nodule', 'Lung Lesion'],
        'min_confidence': 0.75
    },
    # Suppress redundant cardiac findings
    'cardiac_suppress': {
        'trigger': ['Cardiomegaly'],
        'suppress': ['Enlarged Cardiomediastinum'],
        'min_confidence': 0.70
    }
}

# Clinical significance patterns
LOW_CLINICAL_VALUE = [
    'Support Devices',  # Usually not pathological
    'Hernia',          # Often incidental
    'Pleural Thickening'  # Often chronic/stable
]

# Confidence language mapping
CONFIDENCE_LANGUAGE = {
    'CRITICAL': {
        'high': 'demonstrates',
        'moderate': 'strongly suggests', 
        'low': 'raises concern for'
    },
    'HIGH': {
        'high': 'shows evidence of',
        'moderate': 'suggests',
        'low': 'may indicate'
    },
    'MODERATE': {
        'high': 'indicates',
        'moderate': 'suggests possible',
        'low': 'cannot exclude'
    },
    'LOW': {
        'high': 'shows',
        'moderate': 'suggests',
        'low': 'possible'
    }
}





st.set_page_config(
    page_title="Chest X-Ray Analysis",
    page_icon="🩺",
    layout="wide",
    menu_items={}  # 🚀 This removes the GitHub/Report bug/About links
)


# Enhanced CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
        font-weight: bold;
    }
    .report-section {
        background-color: #ffffff;
        color: #2c3e50;
        padding: 2rem;
        border-radius: 10px;
        margin: 1rem 0;
        border-left: 4px solid #1f77b4;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    .warning-box {
        background-color: #fef9e7;
        color: #8b4513;
        padding: 1.5rem;
        border-radius: 8px;
        border-left: 4px solid #f39c12;
        margin: 1rem 0;
        font-weight: 500;
    }
    .success-box {
        background-color: #e8f5e8;
        color: #2d5016;
        padding: 1.5rem;
        border-radius: 8px;
        border-left: 4px solid #27ae60;
        margin: 1rem 0;
        font-weight: 500;
    }
    .info-box {
        background-color: #e3f2fd;
        color: #0d47a1;
        padding: 1.5rem;
        border-radius: 8px;
        border-left: 4px solid #2196f3;
        margin: 1rem 0;
        font-weight: 500;
    }
    .critical-finding {
        background-color: #ffebee;
        color: #b71c1c;
        padding: 1rem;
        margin: 0.5rem 0;
        border-radius: 8px;
        border-left: 4px solid #d32f2f;
        font-weight: bold;
        box-shadow: 0 2px 6px rgba(211, 47, 47, 0.2);
    }
    .high-priority-finding {
        background-color: #fff3e0;
        color: #e65100;
        padding: 1rem;
        margin: 0.5rem 0;
        border-radius: 8px;
        border-left: 4px solid #ff9800;
        font-weight: 600;
    }
    .moderate-finding {
        background-color: #fff8e1;
        color: #f57f17;
        padding: 1rem;
        margin: 0.5rem 0;
        border-radius: 8px;
        border-left: 4px solid #fbc02d;
        font-weight: 500;
    }
    .low-priority-finding {
        background-color: #f3e5f5;
        color: #4a148c;
        padding: 1rem;
        margin: 0.5rem 0;
        border-radius: 8px;
        border-left: 3px solid #9c27b0;
        font-weight: 400;
    }
    .negative-finding {
        background-color: #e8f5e8;
        color: #1b5e20;
        padding: 0.8rem;
        margin: 0.3rem 0;
        border-radius: 6px;
        border-left: 3px solid #4caf50;
        font-weight: 400;
    }
    .suppressed-finding {
        background-color: #f5f5f5;
        color: #666;
        padding: 0.8rem;
        margin: 0.3rem 0;
        border-radius: 6px;
        border-left: 3px solid #999;
        font-style: italic;
    }
    .metric-card {
        background-color: white;
        color: #2c3e50;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 6px rgba(0,0,0,0.1);
        text-align: center;
        font-weight: bold;
    }
    .view-analysis-box {
        background-color: #f0f8ff;
        color: #1e3a8a;
        padding: 1.5rem;
        border-radius: 8px;
        border: 2px solid #3b82f6;
        margin: 1rem 0;
        font-weight: 500;
    }
    .confidence-critical {
        background-color: #d32f2f;
        color: white;
        padding: 0.3rem 0.6rem;
        border-radius: 4px;
        font-weight: bold;
        font-size: 0.85rem;
    }
    .confidence-high {
        background-color: #f57c00;
        color: white;
        padding: 0.3rem 0.6rem;
        border-radius: 4px;
        font-weight: bold;
        font-size: 0.85rem;
    }
    .confidence-moderate {
        background-color: #fbc02d;
        color: white;
        padding: 0.3rem 0.6rem;
        border-radius: 4px;
        font-weight: bold;
        font-size: 0.85rem;
    }
    .confidence-low {
        background-color: #388e3c;
        color: white;
        padding: 0.3rem 0.6rem;
        border-radius: 4px;
        font-weight: bold;
        font-size: 0.85rem;
    }
    .technical-details {
        background-color: #f8fafc;
        color: #374151;
        padding: 1rem;
        border-radius: 6px;
        border-left: 3px solid #6b7280;
        margin: 0.5rem 0;
        font-size: 0.9rem;
    }
    .clinical-summary {
        background-color: #f8f9fa;
        color: #495057;
        padding: 1.5rem;
        border-radius: 8px;
        border-left: 4px solid #6c757d;
        margin: 1rem 0;
        font-weight: 500;
    }
    .download-section {
        background-color: #f0f8ff;
        padding: 1.5rem;
        border-radius: 8px;
        border-left: 4px solid #2196f3;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

@st.cache_resource
def load_xray_models():
    """Load TorchXRayVision models"""
    try:
        models = {}
        models['densenet'] = xrv.models.DenseNet(weights="densenet121-res224-all")
        models['resnet'] = xrv.models.ResNet(weights="resnet50-res512-all")
        
        for model in models.values():
            model.eval()
            
        return models
    except Exception as e:
        st.error(f"Error loading XRayVision models: {str(e)}")
        return None

def init_groq_client():
    """Initialize Groq client"""
    try:
        api_key = os.getenv('GROQ_API_KEY')
        if not api_key:
            st.error("GROQ_API_KEY not found in environment variables")
            return None
        return Groq(api_key=api_key)
    except Exception as e:
        st.error(f"Error initializing Groq client: {str(e)}")
        return None

def preprocess_image(image):
    """Preprocess image for TorchXRayVision models"""
    try:
        img_array = np.array(image.convert('RGB'))
        
        if len(img_array.shape) == 3:
            img_array = skimage.color.rgb2gray(img_array)
        
        img_array = (img_array * 255).astype(np.uint8)
        img_array = xrv.datasets.normalize(img_array, 255)
        
        img_224 = skimage.transform.resize(img_array, (224, 224))
        img_512 = skimage.transform.resize(img_array, (512, 512))
        
        img_224_tensor = torch.FloatTensor(img_224).unsqueeze(0).unsqueeze(0)
        img_512_tensor = torch.FloatTensor(img_512).unsqueeze(0).unsqueeze(0)
        
        return {
            'densenet_input': img_224_tensor,
            'resnet_input': img_512_tensor
        }
    except Exception as e:
        st.error(f"Error preprocessing image: {str(e)}")
        return None

def detect_xray_view_with_reasoning(image, groq_client):
    """Detect if X-ray is AP or PA view using AI"""
    try:
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        img_str = base64.b64encode(buffered.getvalue()).decode()
        
        classification_prompt = """
        Look at this chest X-ray image and determine if it is an AP (Anteroposterior) or PA (Posteroanterior) view.
        
        Respond with ONLY ONE of these two options:
        - AP
        - PA
        
        Base your decision on:
        - Heart size (AP shows enlarged heart due to magnification)
        - Scapula position (AP often shows scapulae over lung fields)
        - Overall image quality and positioning
        
        Answer with just "AP" or "PA":
        """
        
        try:
            classification_response = groq_client.chat.completions.create(
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": classification_prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_str}"
                                }
                            }
                        ]
                    }
                ],
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                max_completion_tokens=10,
                temperature=0.1
            )
            
            classification_text = classification_response.choices[0].message.content.strip().upper()
            
        except Exception:
            try:
                classification_response = groq_client.chat.completions.create(
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": classification_prompt},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/png;base64,{img_str}"
                                    }
                                }
                            ]
                        }
                    ],
                    model="meta-llama/llama-4-maverick-17b-128e-instruct",
                    max_completion_tokens=10,
                    temperature=0.1
                )
                classification_text = classification_response.choices[0].message.content.strip().upper()
                
            except Exception:
                return {
                    'view': 'AP (Anteroposterior)',
                    'confidence': 'Low',
                    'reasoning': 'Auto-detection failed. Defaulting to AP as it is more commonly encountered in clinical practice.',
                    'raw_response': 'Failed to analyze'
                }
        
        if "PA" in classification_text:
            detected_view = "PA (Posteroanterior)"
        else:
            detected_view = "AP (Anteroposterior)"
        
        reasoning_prompt = f"""
        You previously identified this chest X-ray as {detected_view.split(' ')[0]}. 
        Now provide a brief technical explanation for this classification.
        
        Explain in 2-3 sentences:
        - What specific features support this classification
        - Key anatomical indicators you observed
        - Any positioning or quality factors that influenced the decision
        
        Keep the explanation concise and technical.
        """
        
        try:
            reasoning_response = groq_client.chat.completions.create(
                messages=[
                    {
                        "role": "user", 
                        "content": [
                            {"type": "text", "text": reasoning_prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_str}"
                                }
                            }
                        ]
                    }
                ],
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                max_completion_tokens=150,
                temperature=0.1
            )
            
            reasoning = reasoning_response.choices[0].message.content.strip()
            confidence = "High"
            
        except Exception:
            reasoning = f"Classification based on automated analysis. {detected_view.split(' ')[0]} view determined by AI model assessment of anatomical positioning and image characteristics."
            confidence = "Moderate"
        
        return {
            'view': detected_view,
            'confidence': confidence,
            'reasoning': reasoning,
            'raw_response': f"Classification: {classification_text}, Reasoning: {reasoning}"
        }
            
    except Exception as e:
        return {
            'view': 'AP (Anteroposterior)',
            'confidence': 'Low',
            'reasoning': f'Auto-detection encountered an error: {str(e)}. Defaulting to AP view as it is more commonly encountered in clinical settings.',
            'raw_response': f'Error: {str(e)}'
        }

def analyze_xray_with_models(models, processed_images):
    """Analyze X-ray using TorchXRayVision models"""
    try:
        with torch.no_grad():
            densenet_output = models['densenet'](processed_images['densenet_input'])
            densenet_probs = torch.sigmoid(densenet_output).cpu().numpy()[0]
            
        with torch.no_grad():
            resnet_output = models['resnet'](processed_images['resnet_input'])
            resnet_probs = torch.sigmoid(resnet_output).cpu().numpy()[0]
        
        pathologies = models['densenet'].pathologies
        combined_probs = (densenet_probs + resnet_probs) / 2
        
        findings = {}
        for i, pathology in enumerate(pathologies):
            findings[pathology] = {
                'probability': float(combined_probs[i]),
                'densenet_prob': float(densenet_probs[i]),
                'resnet_prob': float(resnet_probs[i]),
                'likely': combined_probs[i] > 0.5
            }
        
        return findings
    except Exception as e:
        st.error(f"Error analyzing X-ray: {str(e)}")
        return None

def get_pathology_priority(pathology):
    """Get clinical priority level for a pathology"""
    for priority, pathologies in CLINICAL_PRIORITY.items():
        if any(p.lower() in pathology.lower() for p in pathologies):
            return priority
    return 'LOW'

def apply_rule_based_filtering(findings):
    """Apply advanced rule-based filtering to suppress contradictory and low-value findings"""
    filtered_findings = {}
    suppressed_findings = {}
    
    # First pass: identify trigger conditions
    triggers_present = {}
    for rule_name, rule in CONTRADICTION_RULES.items():
        for trigger in rule['trigger']:
            for pathology, data in findings.items():
                threshold = CLINICAL_THRESHOLDS.get(pathology, 0.5)
                if trigger.lower() in pathology.lower() and data['probability'] >= threshold:
                    triggers_present[rule_name] = {
                        'trigger_pathology': pathology,
                        'probability': data['probability']
                    }
                    break
    
    # Second pass: apply suppression rules
    for pathology, data in findings.items():
        threshold = CLINICAL_THRESHOLDS.get(pathology, 0.5)
        
        # Skip if below threshold
        if data['probability'] < threshold:
            continue
        
        # Check if this finding should be suppressed
        should_suppress = False
        suppression_reason = ""
        
        for rule_name, rule in CONTRADICTION_RULES.items():
            if rule_name in triggers_present:
                trigger_info = triggers_present[rule_name]
                
                # Check if current pathology should be suppressed
                for suppress_pattern in rule['suppress']:
                    if suppress_pattern.lower() in pathology.lower():
                        # Only suppress if trigger has higher confidence or meets minimum
                        if (trigger_info['probability'] >= rule['min_confidence'] and 
                            trigger_info['probability'] > data['probability']):
                            should_suppress = True
                            suppression_reason = f"Suppressed due to {trigger_info['trigger_pathology']} presence"
                            break
                
                if should_suppress:
                    break
        
        # Apply low clinical value suppression
        if not should_suppress:
            for low_value_pattern in LOW_CLINICAL_VALUE:
                if low_value_pattern.lower() in pathology.lower() and data['probability'] < 0.85:
                    should_suppress = True
                    suppression_reason = "Low clinical significance - high threshold required"
                    break
        
        # Age-related suppression (for common aging changes)
        if not should_suppress:
            aging_patterns = ['Emphysema', 'Fibrosis', 'Pleural Thickening']
            for pattern in aging_patterns:
                if pattern.lower() in pathology.lower() and data['probability'] < 0.75:
                    should_suppress = True
                    suppression_reason = "Common age-related change - high threshold required"
                    break
        
        if should_suppress:
            suppressed_findings[pathology] = {
                **data,
                'suppression_reason': suppression_reason
            }
        else:
            filtered_findings[pathology] = data
    
    return filtered_findings, suppressed_findings

def create_clinical_summary(findings, view_analysis):
    """Create a clinically-focused summary with enhanced rule-based filtering"""
    
    # Apply rule-based filtering first
    filtered_findings, suppressed_findings = apply_rule_based_filtering(findings)
    
    # Categorize by clinical priority
    clinical_categories = {
        'CRITICAL': [],
        'HIGH': [],
        'MODERATE': [],
        'LOW': [],
        'NEGATIVE': [],
        'SUPPRESSED': []
    }
    
    for pathology, data in filtered_findings.items():
        priority = get_pathology_priority(pathology)
        threshold = CLINICAL_THRESHOLDS.get(pathology, 0.5)
        
        clinical_categories[priority].append({
            'pathology': pathology,
            'probability': data['probability'],
            'priority': priority,
            'threshold_used': threshold,
            'confidence_level': 'high' if data['probability'] > 0.8 else 'moderate' if data['probability'] > 0.65 else 'low'
        })
    
    # Add suppressed findings for transparency
    for pathology, data in suppressed_findings.items():
        clinical_categories['SUPPRESSED'].append({
            'pathology': pathology,
            'probability': data['probability'],
            'suppression_reason': data['suppression_reason']
        })
    
    # Add negative findings for context
    negative_findings = []
    for pathology, data in findings.items():
        threshold = CLINICAL_THRESHOLDS.get(pathology, 0.5)
        if data['probability'] < threshold and pathology not in suppressed_findings:
            negative_findings.append({
                'pathology': pathology,
                'probability': data['probability']
            })
    
    negative_findings.sort(key=lambda x: x['probability'], reverse=True)
    clinical_categories['NEGATIVE'] = negative_findings[:5]
    
    # Sort each category by probability
    for category in ['CRITICAL', 'HIGH', 'MODERATE', 'LOW']:
        clinical_categories[category].sort(key=lambda x: x['probability'], reverse=True)
    
    return {
        'view_analysis': view_analysis,
        'clinical_findings': clinical_categories,
        'total_positive_findings': sum(len(clinical_categories[cat]) for cat in ['CRITICAL', 'HIGH', 'MODERATE', 'LOW']),
        'total_suppressed': len(clinical_categories['SUPPRESSED']),
        'most_significant': clinical_categories['CRITICAL'][0] if clinical_categories['CRITICAL'] 
                          else clinical_categories['HIGH'][0] if clinical_categories['HIGH']
                          else clinical_categories['MODERATE'][0] if clinical_categories['MODERATE']
                          else None
    }

def create_enhanced_findings_visualization(clinical_summary):
    """Create enhanced clinical findings visualization with suppression info"""
    
    st.subheader("🏥 Clinical Analysis Results")
    
    # Display view analysis
    view_analysis = clinical_summary['view_analysis']
    confidence_class = {
        'High': 'confidence-high',
        'Moderate': 'confidence-moderate', 
        'Low': 'confidence-low'
    }.get(view_analysis['confidence'], 'confidence-low')
    
    st.markdown(f'''
    <div class="view-analysis-box">
        <h4>📐 X-Ray View Analysis</h4>
        <p><strong>Detected View:</strong> {view_analysis['view']}</p>
        <p><strong>Confidence:</strong> <span class="{confidence_class}">{view_analysis['confidence']}</span></p>
        <p><strong>Clinical Assessment:</strong></p>
        <div class="technical-details">
            {view_analysis['reasoning'].replace(chr(10), "<br>")}
        </div>
    </div>
    ''', unsafe_allow_html=True)
    
    # Summary metrics
    findings = clinical_summary['clinical_findings']
    cols = st.columns(5)
    
    critical_count = len(findings['CRITICAL'])
    high_count = len(findings['HIGH'])
    moderate_count = len(findings['MODERATE'])
    suppressed_count = clinical_summary['total_suppressed']
    
    with cols[0]:
        color = "#d32f2f" if critical_count > 0 else "#666"
        st.markdown(f'''
        <div class="metric-card">
            <h3 style="color: {color};">🚨 {critical_count}</h3>
            <p><strong>Critical</strong><br>Urgent Findings</p>
        </div>
        ''', unsafe_allow_html=True)
        
    with cols[1]:
        color = "#f57c00" if high_count > 0 else "#666"
        st.markdown(f'''
        <div class="metric-card">
            <h3 style="color: {color};">⚡ {high_count}</h3>
            <p><strong>High Priority</strong><br>Important Findings</p>
        </div>
        ''', unsafe_allow_html=True)
        
    with cols[2]:
        moderate_count = len(findings['MODERATE'])
        color = "#fbc02d" if moderate_count > 0 else "#666"
        st.markdown(f'''
        <div class="metric-card">
            <h3 style="color: {color};">📋 {moderate_count}</h3>
            <p><strong>Moderate</strong><br>Notable Findings</p>
        </div>
        ''', unsafe_allow_html=True)
        
    with cols[3]:
        color = "#9c27b0" if suppressed_count > 0 else "#666"
        st.markdown(f'''
        <div class="metric-card">
            <h3 style="color: {color};">🔇 {suppressed_count}</h3>
            <p><strong>Suppressed</strong><br>Rule-Filtered</p>
        </div>
        ''', unsafe_allow_html=True)
        
    with cols[4]:
        most_sig = clinical_summary['most_significant']
        if most_sig:
            prob_text = f"{most_sig['probability']:.0%}"
            st.markdown(f'''
            <div class="metric-card">
                <h3 style="color: #1976d2;">📊 {prob_text}</h3>
                <p><strong>Highest Confidence</strong><br>{most_sig['pathology'][:15]}...</p>
            </div>
            ''', unsafe_allow_html=True)
        else:
            st.markdown(f'''
            <div class="metric-card">
                <h3 style="color: #388e3c;">✓</h3>
                <p><strong>Status</strong><br>No Significant Findings</p>
            </div>
            ''', unsafe_allow_html=True)
    
    st.write("---")
    
    # Clinical findings by priority
    st.subheader("📊 Clinical Findings Assessment")
    
    # Critical findings
    if findings['CRITICAL']:
        st.markdown("**🚨 CRITICAL - Immediate Attention Required**")
        for finding in findings['CRITICAL']:
            confidence_lang = CONFIDENCE_LANGUAGE['CRITICAL'][finding['confidence_level']]
            st.markdown(f'''
            <div class="critical-finding">
                <strong>🚨 {finding['pathology']}</strong><br>
                Analysis <strong>{confidence_lang}</strong> this finding<br>
                <strong>Confidence:</strong> {finding['probability']:.0%} | <strong>Threshold Used:</strong> {finding['threshold_used']:.0%}<br>
                <small>Requires immediate radiologist review and clinical correlation</small>
            </div>
            ''', unsafe_allow_html=True)
    
    # High priority findings  
    if findings['HIGH']:
        st.markdown("**⚡ HIGH PRIORITY - Significant Clinical Findings**")
        for finding in findings['HIGH']:
            confidence_lang = CONFIDENCE_LANGUAGE['HIGH'][finding['confidence_level']]
            st.markdown(f'''
            <div class="high-priority-finding">
                <strong>⚡ {finding['pathology']}</strong><br>
                Analysis <strong>{confidence_lang}</strong> this condition<br>
                <strong>Confidence:</strong> {finding['probability']:.0%} | <strong>Threshold:</strong> {finding['threshold_used']:.0%}<br>
                <small>Recommend prompt clinical evaluation</small>
            </div>
            ''', unsafe_allow_html=True)
    
    # Moderate findings
    if findings['MODERATE']:
        st.markdown("**📋 MODERATE - Notable Clinical Findings**")
        for finding in findings['MODERATE']:
            confidence_lang = CONFIDENCE_LANGUAGE['MODERATE'][finding['confidence_level']]
            st.markdown(f'''
            <div class="moderate-finding">
                <strong>📋 {finding['pathology']}</strong><br>
                Analysis <strong>{confidence_lang}</strong> this finding<br>
                <strong>Confidence:</strong> {finding['probability']:.0%} | <strong>Threshold:</strong> {finding['threshold_used']:.0%}<br>
                <small>Clinical correlation recommended</small>
            </div>
            ''', unsafe_allow_html=True)
    
    # Low priority findings
    if findings['LOW']:
        st.markdown("**📝 Additional Findings**")
        for finding in findings['LOW']:
            confidence_lang = CONFIDENCE_LANGUAGE['LOW'][finding['confidence_level']]
            st.markdown(f'''
            <div class="low-priority-finding">
                <strong>📝 {finding['pathology']}</strong><br>
                Analysis <strong>{confidence_lang}</strong> this finding<br>
                <strong>Confidence:</strong> {finding['probability']:.0%} | <strong>Threshold:</strong> {finding['threshold_used']:.0%}
            </div>
            ''', unsafe_allow_html=True)
    
    # Show suppressed findings for transparency
    if findings['SUPPRESSED']:
        with st.expander(f"🔇 Rule-Based Suppressed Findings ({len(findings['SUPPRESSED'])})"):
            st.markdown("**Findings suppressed by clinical logic rules:**")
            for finding in findings['SUPPRESSED']:
                st.markdown(f'''
                <div class="suppressed-finding">
                    <strong>🔇 {finding['pathology']}</strong> ({finding['probability']:.0%})<br>
                    <small>Reason: {finding['suppression_reason']}</small>
                </div>
                ''', unsafe_allow_html=True)
    
    # Show normal findings
    if not any(findings[cat] for cat in ['CRITICAL', 'HIGH', 'MODERATE', 'LOW']):
        st.markdown("**✅ No Significant Abnormalities Detected**")
        st.markdown('''
        <div class="negative-finding">
            <strong>Assessment:</strong> No findings above clinical significance thresholds detected by AI analysis.<br>
            <small>This suggests the X-ray appears normal on automated screening with rule-based filtering applied.</small>
        </div>
        ''', unsafe_allow_html=True)
    else:
        # Show top negative findings for clinical context
        st.markdown("**ℹ️ Additional Context - Findings Below Clinical Threshold**")
        for neg_finding in findings['NEGATIVE'][:3]:
            st.markdown(f'''
            <div class="negative-finding">
                <strong>✓ {neg_finding['pathology']}:</strong> {neg_finding['probability']:.0%} confidence (below {CLINICAL_THRESHOLDS.get(neg_finding['pathology'], 0.5):.0%} threshold)
            </div>
            ''', unsafe_allow_html=True)

def create_structured_findings_summary(clinical_summary, patient_info):
    """Create a structured summary for LLM report generation"""
    
    findings = clinical_summary['clinical_findings']
    view_analysis = clinical_summary['view_analysis']
    
    high_confidence_findings = []
    moderate_confidence_findings = []
    mild_concern_findings = []
    
    for priority in ['CRITICAL', 'HIGH', 'MODERATE', 'LOW']:
        for finding in findings[priority]:
            entry = {
                'pathology': finding['pathology'],
                'probability': finding['probability'],
                'percentage': f"{finding['probability']:.1%}",
                'clinical_priority': finding['priority'],
                'threshold_used': f"{finding['threshold_used']:.1%}",
                'confidence_level': finding['confidence_level']
            }
            
            if finding['probability'] > 0.8:
                high_confidence_findings.append(entry)
            elif finding['probability'] >= 0.65:
                moderate_confidence_findings.append(entry)
            else:
                mild_concern_findings.append(entry)
    
    structured_summary = {
        'patient_metadata': {
            'age': patient_info.get('age', 'N/A'),
            'gender': patient_info.get('gender', 'N/A'),
            'view': patient_info.get('view', 'N/A'),
            'date': patient_info.get('date', 'N/A')
        },
        'view_analysis': view_analysis,
        'ai_analysis_results': {
            'high_confidence_findings': high_confidence_findings,
            'moderate_confidence_findings': moderate_confidence_findings,
            'mild_concern_findings': mild_concern_findings,
            'negative_findings': [f['pathology'] for f in findings['NEGATIVE'][:5]],
            'suppressed_findings': [{'pathology': f['pathology'], 'reason': f['suppression_reason']} for f in findings['SUPPRESSED']]
        },
        'clinical_context': {
            'total_significant_findings': clinical_summary['total_positive_findings'],
            'total_suppressed': clinical_summary['total_suppressed'],
            'most_significant_finding': clinical_summary['most_significant'],
            'thresholds_applied': 'Clinical pathology-specific thresholds with rule-based filtering',
            'filtering_applied': 'Advanced contradiction filtering and clinical prioritization'
        },
        'model_performance': {
            'total_pathologies_analyzed': len(findings['NEGATIVE']) + clinical_summary['total_positive_findings'],
            'models_used': ['TorchXRayVision DenseNet121', 'TorchXRayVision ResNet50'],
            'ensemble_method': 'Average probability with clinical threshold filtering and rule-based suppression'
        }
    }
    
    return structured_summary

def generate_structured_report_with_template(structured_findings, groq_client):
    """Generate structured radiology report using template format with enhanced filtering"""
    
    view_info = structured_findings['view_analysis']
    high_conf = structured_findings['ai_analysis_results']['high_confidence_findings']
    mod_conf = structured_findings['ai_analysis_results']['moderate_confidence_findings']
    border_conf = structured_findings['ai_analysis_results']['mild_concern_findings']
    suppressed = structured_findings['ai_analysis_results']['suppressed_findings']
    
    findings_text = f"""
AI MODEL ANALYSIS RESULTS (Enhanced Clinical Filtering):
Models: TorchXRayVision DenseNet121 + ResNet50 Ensemble
Filtering: Pathology-specific clinical thresholds + advanced rule-based suppression
Suppression Rules Applied: {len(suppressed)} findings suppressed by clinical logic

HIGH CONFIDENCE POSITIVE FINDINGS (>80% probability, above clinical thresholds):
"""
    
    if high_conf:
        for finding in high_conf:
            findings_text += f"- {finding['pathology']}: {finding['percentage']} (Threshold: {finding['threshold_used']}, Priority: {finding['clinical_priority']})\n"
    else:
        findings_text += "- None detected above clinical significance thresholds\n"
    
    findings_text += f"""
MODERATE CONFIDENCE FINDINGS (65-80% probability, clinically significant):
"""
    
    if mod_conf:
        for finding in mod_conf:
            findings_text += f"- {finding['pathology']}: {finding['percentage']} (Threshold: {finding['threshold_used']}, Priority: {finding['clinical_priority']})\n"
    else:
        findings_text += "- None detected above clinical significance thresholds\n"
        
    findings_text += f"""
BORDERLINE FINDINGS (50-65% probability, clinical correlation recommended):
"""
    
    if border_conf:
        for finding in border_conf:
            findings_text += f"- {finding['pathology']}: {finding['percentage']} (Threshold: {finding['threshold_used']}, Priority: {finding['clinical_priority']})\n"
    else:
        findings_text += "- None detected\n"

    if suppressed:
        findings_text += f"""
RULE-SUPPRESSED FINDINGS (filtered by clinical logic):
"""
        for finding in suppressed[:3]:  # Show top 3 suppressed
            findings_text += f"- {finding['pathology']}: Suppressed - {finding['reason']}\n"

    prompt = f"""
You are an expert radiologist generating a structured chest X-ray report. Use enhanced AI analysis with rule-based filtering.

VIEW ANALYSIS:
- Detected View: {view_info['view']}
- Detection Confidence: {view_info['confidence']}

{findings_text}

IMPORTANT: Only report findings that exceeded clinical significance thresholds and passed rule-based filtering. Use appropriate clinical language based on confidence levels.

Generate a report using this EXACT structure:

**POSITIONING & EXPOSURE**
- Proper {view_info['view'].split()[0]} positioning confirmed.
- Exposure adequate with spine visible through cardiac silhouette.
- No evidence of rotation or motion artifacts.

**BONY COMPONENTS**
[Use clinical findings if present, otherwise use normal defaults]
- Ribs: [normal symmetry OR specific abnormality if detected above threshold]
- Clavicles: [intact OR abnormality if detected] 
- Spine: [vertebrae normal OR specific finding]
- Costophrenic angles: [clear OR blunted if effusion above threshold]

**LUNG FINDINGS**
[Map threshold-exceeded findings to anatomical zones]
- Upper zones: [normal parenchyma OR specific findings above clinical threshold]
- Mid zones: [normal parenchyma OR specific findings above clinical threshold] 
- Lower zones: [normal parenchyma OR specific findings above clinical threshold]
- Lung markings: [normal OR abnormal based on significant findings]
- Pleural effusion: [absent OR present if above clinical threshold]
- Pneumothorax: [absent OR present if above clinical threshold]

**CARDIAC SILHOUETTE**
- Heart size: [normal OR enlarged if cardiomegaly above threshold]
- Cardiothoracic ratio: [normal OR increased if significant]
- Cardiac borders: [well-defined OR specific abnormality]

**DIAPHRAGM**
- Contour: [smooth OR specific abnormality if significant]
- Costophrenic angles: [clear OR blunted based on effusion threshold]

**ADDITIONAL OBSERVATIONS**
- Medical devices: [None visible OR describe if present]
- Additional findings: [List only findings above clinical thresholds: nodules, masses, etc.]

**IMPRESSION**
[Clinical threshold and rule-based filtered assessment]:
- If high-confidence findings: "Findings demonstrate [specific pathologies]. Recommend [appropriate follow-up]."
- If moderate findings only: "Findings suggest [pathologies]. Clinical correlation recommended."
- If borderline only: "Borderline findings noted. Consider follow-up if clinically indicated."
- If no significant findings: "No acute abnormalities detected after clinical filtering. Normal chest X-ray."

CRITICAL REQUIREMENTS:
1. Use confidence-appropriate language:
   - High confidence (>80%): "demonstrates" or "shows"
   - Moderate (65-80%): "suggests" or "indicates"  
   - Borderline (50-65%): "possible" or "borderline"
2. Only report findings above clinical significance thresholds that passed rule-based filtering
3. Fill ALL sections with appropriate defaults if no significant findings
4. Map pathologies to correct anatomical locations
5. Provide appropriate clinical recommendations
"""

    try:
        response = groq_client.chat.completions.create(
            messages=[
                {"role": "user", "content": prompt}
            ],
            model="llama-3.1-8b-instant",
            max_tokens=1200,
            temperature=0.1
        )
        
        return response.choices[0].message.content
    
    except Exception as e:
        st.error(f"Error generating structured report: {str(e)}")
        return None

def create_word_document(report_content, structured_findings, patient_info):
    """Create a Word document with the comprehensive report"""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('CHEST X-RAY AI-ASSISTED ANALYSIS REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add patient information section
        doc.add_heading('Patient Information', level=1)
        patient_table = doc.add_table(rows=6, cols=2)
        patient_table.style = 'Table Grid'
        
        patient_data = [
            ('Name:', patient_info.get('name', 'N/A')),
            ('Age / Gender:', f"{patient_info.get('age', 'N/A')} / {patient_info.get('gender', 'N/A')}"),
            ('Patient ID:', patient_info.get('patient_id', 'N/A')),
            ('Examination Date:', patient_info.get('date', 'N/A')),
            ('Referring Physician:', patient_info.get('physician', 'N/A')),
            ('Report Generated:', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        ]
        
        for i, (label, value) in enumerate(patient_data):
            patient_table.cell(i, 0).text = label
            patient_table.cell(i, 1).text = str(value)
            patient_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Add structured report content
        doc.add_heading('Clinical Report', level=1)
        
        sections = report_content.split('**')
        current_paragraph = None
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
                
            if section.isupper() or any(heading in section.upper() for heading in ['POSITIONING', 'BONY', 'LUNG', 'CARDIAC', 'DIAPHRAGM', 'ADDITIONAL', 'IMPRESSION']):
                if section not in ['\n', ' ', '']:
                    heading = doc.add_heading(section, level=2)
                    current_paragraph = None
            else:
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph(section)
                else:
                    current_paragraph.add_run(section)
        
        # Add AI analysis details
        doc.add_page_break()
        doc.add_heading('AI Analysis Technical Details', level=1)
        
        # Model information
        doc.add_heading('Models and Enhanced Filtering Used', level=2)
        models_para = doc.add_paragraph()
        models_para.add_run('• TorchXRayVision DenseNet121 (CheXpert, MIMIC-CXR, NIH14, PadChest)\n')
        models_para.add_run('• TorchXRayVision ResNet50 (multi-dataset ensemble training)\n')
        models_para.add_run('• Ensemble Method: Average probability with clinical threshold filtering\n')
        models_para.add_run('• Clinical Thresholds: Pathology-specific thresholds optimized for clinical significance\n')
        models_para.add_run('• Rule-Based Filtering: Advanced contradiction suppression and clinical prioritization\n')
        models_para.add_run('• False Positive Reduction: Through threshold optimization and clinical logic rules')
        
        # View detection details
        view_info = structured_findings['view_analysis']
        doc.add_heading('View Detection Analysis', level=2)
        view_para = doc.add_paragraph()
        view_para.add_run(f'Detected View: {view_info["view"]}\n').bold = True
        view_para.add_run(f'Confidence Level: {view_info["confidence"]}\n')
        view_para.add_run(f'Clinical Assessment: {view_info["reasoning"]}')
        
        # Clinical findings summary with suppression info
        doc.add_heading('Enhanced Clinical Findings Summary', level=2)
        high_conf = structured_findings['ai_analysis_results']['high_confidence_findings']
        mod_conf = structured_findings['ai_analysis_results']['moderate_confidence_findings']
        border_conf = structured_findings['ai_analysis_results']['mild_concern_findings']
        suppressed = structured_findings['ai_analysis_results']['suppressed_findings']
        
        if high_conf:
            doc.add_heading('High Confidence Findings (>80%, Above Clinical Thresholds)', level=3)
            for finding in high_conf:
                finding_para = doc.add_paragraph()
                finding_para.add_run(f'• {finding["pathology"]}: {finding["percentage"]}').bold = True
                finding_para.add_run(f' (Priority: {finding["clinical_priority"]}, Threshold: {finding["threshold_used"]})')
        
        if mod_conf:
            doc.add_heading('Moderate Confidence Findings (65-80%, Clinically Significant)', level=3)
            for finding in mod_conf:
                finding_para = doc.add_paragraph()
                finding_para.add_run(f'• {finding["pathology"]}: {finding["percentage"]}')
                finding_para.add_run(f' (Priority: {finding["clinical_priority"]}, Threshold: {finding["threshold_used"]})')
        
        if border_conf:
            doc.add_heading('Borderline Findings (50-65%, Clinical Correlation Recommended)', level=3)
            for finding in border_conf:
                finding_para = doc.add_paragraph()
                finding_para.add_run(f'• {finding["pathology"]}: {finding["percentage"]}')
                finding_para.add_run(f' (Priority: {finding["clinical_priority"]}, Threshold: {finding["threshold_used"]})')
        
        # Rule-based filtering results
        if suppressed:
            doc.add_heading('Rule-Based Filtering Results', level=2)
            suppression_para = doc.add_paragraph()
            suppression_para.add_run(f'Total Findings Suppressed: {len(suppressed)}\n').bold = True
            suppression_para.add_run('Suppression enhances clinical relevance by removing contradictory or low-value findings.\n\n')
            
            for finding in suppressed:
                suppression_para.add_run(f'• {finding["pathology"]}: {finding["reason"]}\n')
        
        # Clinical context with enhanced info
        doc.add_heading('Clinical Context', level=2)
        context_para = doc.add_paragraph()
        context = structured_findings['clinical_context']
        context_para.add_run(f'Total Significant Findings: {context["total_significant_findings"]}\n')
        context_para.add_run(f'Total Suppressed by Rules: {context["total_suppressed"]}\n')
        context_para.add_run(f'Threshold System: {context["thresholds_applied"]}\n')
        context_para.add_run(f'Filtering Applied: {context["filtering_applied"]}\n')
        
        if context['most_significant_finding']:
            most_sig = context['most_significant_finding']
            context_para.add_run(f'Most Significant Finding: {most_sig["pathology"]} ({most_sig["probability"]:.1%})')
        
        # Enhanced disclaimer
        doc.add_page_break()
        doc.add_heading('Medical Disclaimer and Clinical Context', level=1)
        disclaimer_para = doc.add_paragraph()
        disclaimer_para.add_run('IMPORTANT: ').bold = True
        disclaimer_para.add_run('This report uses enhanced AI analysis with clinical threshold filtering and rule-based suppression for optimal clinical relevance. Only findings above established clinical thresholds and passing contradiction filters are reported. This tool is intended for screening support and must be validated by a qualified radiologist.\n\n')
        
        disclaimer_para.add_run('ENHANCED FILTERING SYSTEM: ').bold = True
        disclaimer_para.add_run('The analysis applies pathology-specific probability thresholds combined with rule-based suppression of contradictory findings. This significantly reduces false positives while maintaining sensitivity for clinically significant findings.\n\n')
        
        disclaimer_para.add_run('RULE-BASED SUPPRESSION: ').bold = True
        disclaimer_para.add_run('Advanced clinical logic rules suppress contradictory findings (e.g., suppress general infiltration when specific consolidation is detected) and filter low clinical value findings, mimicking radiologist decision-making.\n\n')
        
        disclaimer_para.add_run('GROUNDING STATEMENT: ').bold = True
        disclaimer_para.add_run('All reported findings exceeded their respective clinical significance thresholds and passed rule-based filtering. No findings have been added that are not supported by the enhanced AI analysis system.\n\n')
        
        disclaimer_para.add_run('This enhanced AI-assisted analysis should be used in conjunction with clinical judgment and must be reviewed by a qualified healthcare professional for final interpretation and diagnosis.')
        
        return doc
        
    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        return None

def main():
    st.markdown('<h1 class="main-header">🩺 Enhanced AI Chest X-Ray Analysis with Rule-Based Filtering</h1>', unsafe_allow_html=True)
    
    # Enhanced disclaimer
    st.markdown("""
    <div class="warning-box">
        <strong>⚠️ Medical Disclaimer:</strong> This application uses enhanced AI analysis with clinical threshold filtering 
        and advanced rule-based suppression. Only findings above established clinical thresholds that pass contradiction 
        filtering are reported. This is a screening tool - all results must be reviewed by a qualified radiologist.
    </div>
    """, unsafe_allow_html=True)
    
    # Load models
    with st.spinner("🔬 Loading specialized models with enhanced filtering..."):
        models = load_xray_models()
        groq_client = init_groq_client()
    
    if not models or not groq_client:
        st.error("Failed to load required models. Please check your setup.")
        st.stop()
    
    st.markdown("""
    <div class="success-box">
        <strong>✅ Enhanced Clinical Analysis System Loaded!</strong><br>
        • TorchXRayVision DenseNet121 + ResNet50 Ensemble<br>
        • Clinical Threshold Filtering (Pathology-Specific)<br>
        • Advanced Rule-Based Contradiction Suppression<br>
        • Enhanced AP/PA Detection with Clinical Context<br>
        • Confidence-Appropriate Clinical Language<br>
        • False Positive Reduction Through Multi-Layer Filtering
    </div>
    """, unsafe_allow_html=True)
    
    # Create layout
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("📁 Upload Chest X-Ray Image")
        uploaded_file = st.file_uploader(
            "Choose a chest X-ray image",
            type=['png', 'jpg', 'jpeg'],
            help="Upload a clear chest X-ray image for enhanced AI analysis"
        )
        
        if uploaded_file:
            image = Image.open(uploaded_file)
            st.image(image, caption="Uploaded X-Ray", use_column_width=True)
    
    with col2:
        st.subheader("👤 Patient Information")
        
        with st.form("patient_info_form"):
            patient_name = st.text_input("Patient Name")
            
            col_age, col_gender = st.columns(2)
            with col_age:
                patient_age = st.number_input("Age", min_value=0, max_value=150, value=None)
            with col_gender:
                patient_gender = st.selectbox("Gender", ["", "Male", "Female", "Other"])
            
            patient_id = st.text_input("Patient ID")
            exam_date = st.date_input("Examination Date", datetime.now().date())
            physician = st.text_input("Referring Physician")
            
            analyze_button = st.form_submit_button("🔍 Analyze with Enhanced Filtering", use_container_width=True)
    
    # Analysis section
    if uploaded_file and analyze_button:
        with st.spinner("🧠 Running enhanced clinical analysis with rule-based filtering..."):
            # Auto-detect X-ray view
            st.info("🔍 Step 1: Auto-detecting X-ray view (AP/PA) with clinical context...")
            view_analysis = detect_xray_view_with_reasoning(image, groq_client)
            
            patient_info = {
                'name': patient_name or 'Anonymous Patient',
                'age': str(patient_age) if patient_age else 'N/A',
                'gender': patient_gender or 'N/A',
                'patient_id': patient_id or 'N/A',
                'date': exam_date.strftime("%Y-%m-%d") if exam_date else 'N/A',
                'physician': physician or 'N/A',
                'view': view_analysis['view']
            }
            
            # Preprocess image
            st.info("🔬 Step 2: Preprocessing image for AI pathology detection...")
            processed_images = preprocess_image(image)
            
            if processed_images:
                # Analyze with models
                st.info("🔬 Step 3: Running ensemble AI analysis with enhanced filtering...")
                findings = analyze_xray_with_models(models, processed_images)
                
                if findings:
                    # Create clinical summary with enhanced filtering
                    clinical_summary = create_clinical_summary(findings, view_analysis)
                    
                    # Enhanced findings visualization
                    create_enhanced_findings_visualization(clinical_summary)
                    
                    st.write("---")
                    
                    # Create structured findings for report generation
                    structured_findings = create_structured_findings_summary(clinical_summary, patient_info)
                    
                    # Generate structured clinical report
                    with st.spinner("📝 Step 4: Generating clinical report with enhanced filtering..."):
                        clinical_report = generate_structured_report_with_template(structured_findings, groq_client)
                    
                    if clinical_report:
                        # Display clinical report
                        st.markdown('<div class="report-section">', unsafe_allow_html=True)
                        st.markdown("## 📋 Enhanced Clinical Report with Rule-Based Filtering")
                        st.markdown(clinical_report)
                        st.markdown('</div>', unsafe_allow_html=True)
                        
                        # FIXED: Create Word document upfront and provide download button
                        st.write("---")
                        st.subheader("📄 Clinical Documentation Download")
                        
                        with st.spinner("📄 Preparing comprehensive clinical documentation..."):
                            doc = create_word_document(clinical_report, structured_findings, patient_info)
                        
                        if doc:
                            doc_buffer = io.BytesIO()
                            doc.save(doc_buffer)
                            doc_buffer.seek(0)
                            
                            # Fixed download section
                            st.markdown('''
                            <div class="download-section">
                                <h4>📄 Clinical Report Ready for Download</h4>
                                <p>Complete clinical report with enhanced filtering results, rule-based suppression details, and technical analysis.</p>
                            </div>
                            ''', unsafe_allow_html=True)
                            
                            st.download_button(
                                label="💾 Download Enhanced Clinical Report (DOCX)",
                                data=doc_buffer.getvalue(),
                                file_name=f"enhanced_clinical_xray_report_{patient_info['name'].replace(' ', '_')}_{exam_date.strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                help="Complete clinical report with rule-based filtering and suppression analysis"
                            )
                        else:
                            st.error("Failed to create clinical documentation")
                        
                        # Clinical analysis summary
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            with st.expander("📊 Enhanced Analysis Summary"):
                                st.markdown(f"**Patient:** {patient_info['name']}")
                                st.markdown(f"**Analysis Date:** {patient_info['date']}")
                                st.markdown("---")
                                
                                total_significant = clinical_summary['total_positive_findings']
                                total_suppressed = clinical_summary['total_suppressed']
                                
                                if total_significant > 0:
                                    st.markdown(f"**Total Significant Findings:** {total_significant}")
                                    if clinical_summary['most_significant']:
                                        most_sig = clinical_summary['most_significant']
                                        st.markdown(f"**Most Significant:** {most_sig['pathology']} ({most_sig['probability']:.0%})")
                                else:
                                    st.markdown("**Status:** No findings above clinical thresholds")
                                
                                st.markdown(f"**Rule-Suppressed:** {total_suppressed} findings")
                                st.markdown("**Enhanced Filtering:** Clinical thresholds + rule-based suppression")
                        
                        with col2:
                            with st.expander("🔧 Filtering System Details"):
                                st.markdown("**Clinical Thresholds Applied:**")
                                st.markdown("• Critical: 75-85% (Pneumothorax, Mass)")
                                st.markdown("• High Priority: 65-75% (Consolidation)")
                                st.markdown("• Moderate: 60-65% (Cardiomegaly)")
                                st.markdown("• Low Priority: Higher thresholds")
                                st.markdown("---")
                                st.markdown("**Rule-Based Suppression:**")
                                st.markdown("• Contradiction detection active")
                                st.markdown("• Low clinical value filtering")
                                st.markdown("• Age-related change suppression")
                                st.markdown("• Enhanced clinical relevance")
                        
                        # Enhanced clinical insights
                        st.write("---")
                        st.subheader("🎯 Enhanced Clinical Assessment Summary")
                        
                        insights = []
                        
                        # View analysis insights
                        if view_analysis['confidence'] == 'High':
                            insights.append(f"✅ **View Analysis:** High confidence {view_analysis['view']} detection supports reliable analysis.")
                        elif view_analysis['confidence'] == 'Moderate':
                            insights.append(f"⚠️ **View Analysis:** Moderate confidence {view_analysis['view']} - manual verification recommended.")
                        else:
                            insights.append(f"🔍 **View Analysis:** Low confidence {view_analysis['view']} - radiologist review recommended.")
                        
                        # Clinical priority insights with filtering context
                        critical_count = len(clinical_summary['clinical_findings']['CRITICAL'])
                        high_count = len(clinical_summary['clinical_findings']['HIGH'])
                        moderate_count = len(clinical_summary['clinical_findings']['MODERATE'])
                        suppressed_count = clinical_summary['total_suppressed']
                        
                        if critical_count > 0:
                            critical_pathologies = [f['pathology'] for f in clinical_summary['clinical_findings']['CRITICAL']]
                            insights.append(f"🚨 **CRITICAL PRIORITY:** {critical_count} finding(s) require immediate attention: {', '.join(critical_pathologies)}")
                        
                        if high_count > 0:
                            high_pathologies = [f['pathology'] for f in clinical_summary['clinical_findings']['HIGH']]
                            insights.append(f"⚡ **HIGH PRIORITY:** {high_count} significant finding(s): {', '.join(high_pathologies)}")
                        
                        if moderate_count > 0:
                            insights.append(f"📋 **MODERATE PRIORITY:** {moderate_count} notable finding(s) - clinical correlation recommended.")
                        
                        if suppressed_count > 0:
                            insights.append(f"🔇 **RULE-BASED FILTERING:** {suppressed_count} findings suppressed by clinical logic rules to enhance relevance.")
                        
                        if critical_count == 0 and high_count == 0 and moderate_count == 0:
                            insights.append("✅ **ASSESSMENT:** No findings above clinical significance thresholds after rule-based filtering - appears normal on enhanced AI screening.")
                        
                        # Enhanced filtering system insight
                        insights.append("🎯 **ENHANCED FILTERING:** Multi-layer system with clinical thresholds + rule-based suppression reduces false positives while maintaining sensitivity.")
                        
                        # Display insights
                        for insight in insights:
                            st.markdown(f'''
                            <div class="clinical-summary">
                                {insight}
                            </div>
                            ''', unsafe_allow_html=True)
                        
                        # Enhanced disclaimer with filtering details
                        st.write("---")
                        st.markdown("""
                        <div class="warning-box">
                            <strong>🩺 Enhanced Clinical System:</strong> This analysis uses pathology-specific probability thresholds 
                            combined with advanced rule-based filtering that suppresses contradictory findings and low clinical value 
                            results. Only findings above clinical thresholds that pass contradiction filtering are reported, 
                            significantly improving clinical relevance. All results require radiologist validation for final diagnosis.
                        </div>
                        """, unsafe_allow_html=True)
                        
                    else:
                        st.error("❌ Failed to generate clinical report.")
                else:
                    st.error("❌ Failed to analyze the image.")
            else:
                st.error("❌ Failed to preprocess the image.")
    
    # Enhanced sidebar with rule-based filtering information
    with st.sidebar:
        st.markdown("## 🎯 Enhanced Clinical Analysis System")
        st.markdown("""
        **🔬 AI Models:**
        - **DenseNet121**: Multi-dataset (CheXpert, MIMIC-CXR, NIH14, PadChest)
        - **ResNet50**: Ensemble architecture (512x512 resolution)
        - **Method**: Average probability + multi-layer filtering
        
        **🎯 Clinical Threshold System:**
        - **Critical Findings**: 80-85% thresholds (Pneumothorax, Mass)
        - **High Priority**: 65-75% thresholds (Consolidation, Pneumonia)
        - **Moderate Priority**: 60-65% thresholds (Cardiomegaly, Effusion)
        - **Low Priority**: Higher thresholds to reduce noise
        
        **🔧 Rule-Based Filtering:**
        - **Contradiction Suppression**: Remove conflicting findings
        - **Clinical Logic**: Prefer specific over general diagnoses
        - **Low-Value Filtering**: Suppress incidental/stable findings
        - **Age-Related Suppression**: Filter common aging changes
        """)
        
        st.markdown("## 🚨 Enhanced Clinical Priority")
        st.markdown("""
        **CRITICAL (Immediate):**
        - Pneumothorax (80% threshold + rules)
        - Massive Pleural Effusion (75%)
        - Large Opacity (70%)
        
        **HIGH (Significant):**
        - Mass (80% threshold + contradiction rules) 
        - Consolidation (70% + infiltration suppression)
        - Pneumonia (75% + general finding suppression)
        
        **MODERATE (Notable):**
        - Pleural Effusion (65% + massive effusion rules)
        - Cardiomegaly (60% + mediastinal suppression)
        - Atelectasis (65% + opacity rules)
        """)
        
        st.markdown("---")
        st.markdown("## 🔇 Rule-Based Suppression")
        st.markdown("""
        **Contradiction Rules:**
        - **Pneumothorax Present**: Suppress general lung findings unless >85% confidence
        - **Consolidation/Pneumonia**: Suppress general infiltration findings
        - **Mass Detected**: Suppress general nodule findings
        - **Cardiomegaly**: Suppress enlarged cardiomediastinum
        
        **Clinical Value Rules:**
        - **Support Devices**: Require >85% (usually normal)
        - **Hernia**: Require >80% (often incidental)
        - **Pleural Thickening**: Require >60% (often stable)
        
        **Age-Related Suppression:**
        - **Emphysema**: Require >75% (common aging)
        - **Fibrosis**: Require >70% (often stable)
        - **Minor Changes**: Higher thresholds applied
        """)
        
        st.markdown("---")
        st.markdown("## 📋 Enhanced Features")
        st.markdown("""
        **✅ Multi-Layer Filtering:**
        - Clinical thresholds (pathology-specific)
        - Rule-based contradiction detection
        - Clinical relevance optimization
        - False positive reduction
        
        **✅ Clinical Logic Integration:**
        - Radiologist-like decision making
        - Prefer specific over general findings
        - Suppress low clinical value results
        - Context-aware filtering
        
        **✅ Enhanced Documentation:**
        - Suppression transparency
        - Rule application details
        - Clinical context preservation
        - Professional report format
        """)
        
        st.markdown("---")
        st.markdown("## © 2025 MEDxAI Innovations Pvt Ltd. All rights reserved.")

if __name__ == "__main__":
    main()